#!/usr/bin/env python3
"""Generate a polished Word (.docx) book proposal.

Usage:
  python3 build_book_proposal_docx.py [OUT_DOCX]

Defaults to:
  /Users/arthur/clawd/ComputationalPathsLean/Computational_Paths_Book_Proposal_CUP.docx

Notes:
- Pure mathematics proposal; Lean appears only in Appendix C (methodology).
- Chapter 1 outline expanded per instructions (20–25 pages inside the eventual book).
"""

from __future__ import annotations

import sys
import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DEFAULT_OUT = Path(
    "/Users/arthur/clawd/ComputationalPathsLean/Computational_Paths_Book_Proposal_CUP.docx"
)

TITLE = "Computational Paths and the Algebraic Topology of Rewriting"
SUBTITLE = "From Labelled Deduction to Higher Groupoids"


def set_default_font(doc: Document, name: str = "Times New Roman", size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = name
    font.size = Pt(size_pt)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    r.italic = True
    r.font.size = Pt(14)

    doc.add_paragraph().add_run("\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Book Proposal")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Authors: Ruy de Queiroz; Arthur Ramos (to be confirmed)")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Prepared: {date.today().strftime('%B %Y')}")

    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def main(out_path: Path, template_path: Path | None = None) -> None:
    # If a publisher template is provided, load it to preserve its styles/headers.
    doc = Document(str(template_path)) if template_path else Document()
    set_default_font(doc)

    # If the template already contains meaningful content, append after a page break.
    if template_path:
        nonempty = [p for p in doc.paragraphs if p.text.strip()]
        if nonempty:
            doc.add_page_break()

    add_title_page(doc)

    # Executive summary / overview
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This book develops a mathematical theory in which the algebra of rewriting—"
        "as it appears in labelled natural deduction—gives rise to groupoids and higher"
        " groupoids. The central thesis is that rewrite traces are intrinsically homotopical:"
        " when equality proofs are equipped with explicit reasons (rewrite sequences), one"
        " obtains a structured ‘path algebra’ whose higher cells arise from rewriting between"
        " rewrites."
    )
    doc.add_paragraph(
        "The book is written as pure mathematics at the interface of proof theory, term rewriting,"
        " higher-dimensional algebra, and algebraic topology. Formal verification in Lean is mentioned"
        " only as a methodological appendix; no prior knowledge of proof assistants is assumed."
    )

    doc.add_paragraph("Core technical deliverables include:")
    add_bullets(
        doc,
        [
            "A groupoid of computational paths where composition is transitivity and inversion is symmetry, with laws derived from a terminating, confluent rewrite system.",
            "A hierarchy of higher cells (2-cells as rewrite derivations between paths, higher coherences thereafter) yielding a weak ω-groupoid structure.",
            "A derivation of path induction (the J-eliminator) as a theorem from contractibility of the based path space (within the computational-paths framework).",
            "An explicit treatment of the two-level architecture Path vs Eq: proof-relevant rewrite traces refining proof-irrelevant propositional equality, addressing UIP upfront.",
            "Applications to fundamental groups/groupoids and classical calculations (circle, torus, Klein bottle, projective plane), including Seifert–van Kampen.",
        ],
    )

    doc.add_paragraph(
        "Estimated length (target for publication): 400–500 pages, including appendices."
    )

    # Positioning / architectural honesty (from critical analysis)
    doc.add_heading("Scope and Architectural Choice (Path vs Eq)", level=1)
    doc.add_paragraph(
        "A key design choice is to distinguish: (i) ordinary propositional equality (Eq), which is proof-irrelevant"
        " in classical foundations; and (ii) computational paths (Path), which remember the rewrite trace witnessing"
        " an identification. Path is not a new foundation competing with HoTT or cubical type theory; rather, it is"
        " a mathematically motivated refinement that records explicit reasons for equality and supports a higher"
        " algebraic structure derived from rewriting. The proposal treats this two-level architecture explicitly and"
        " early, so referees can see precisely what is captured and what is not."
    )

    # Narrative arc
    doc.add_heading("The Narrative Arc (Three Acts)", level=1)
    doc.add_heading("Act I — The algebra of proof-terms", level=2)
    doc.add_paragraph(
        "Beginning with labelled natural deduction (LND), we treat ‘reasons for equality’ (rewrite traces) as first-class"
        " mathematical objects. We develop the operations on traces (reflexivity, symmetry, transitivity, congruence,"
        " transport) and study a terminating and confluent rewrite system governing their algebra."
    )
    doc.add_heading("Act II — Higher structure emerges", level=2)
    doc.add_paragraph(
        "Rewrite derivations between paths become 2-cells; higher rewrites yield higher cells. Iterating, we obtain a tower"
        " of globular data forming a weak ω-groupoid in the sense of Batanin/Leinster, with explicit coherence witnesses"
        " arising from the rewrite system."
    )
    doc.add_heading("Act III — Connections and applications", level=2)
    doc.add_paragraph(
        "We derive J, clarify the failure and role of UIP at the Path-level, compare with HoTT and higher-dimensional rewriting"
        " (Squier, Guiraud–Malbos), and compute fundamental groups/groupoids of standard spaces in a rewriting-driven framework."
    )

    # Chapter outline
    doc.add_heading("Chapter-by-Chapter Outline", level=1)

    # Part I
    doc.add_heading("Part I — Foundations: The Algebra of Computational Paths", level=1)

    doc.add_heading("Chapter 1 — Introduction and Motivation (expanded; 20–25 pages)", level=2)
    doc.add_paragraph(
        "Chapter 1 is expanded to situate the project as the mathematical continuation of Ruy de Queiroz’s programme,"
        " moving from labelled deduction and rewrite reasons to higher groupoids. It includes a detailed account of the"
        " historical and conceptual steps leading to computational paths, and makes the Path vs Eq architecture explicit."
    )

    doc.add_heading("§1.1 The Curry–Howard correspondence and its discontents (≈3 pages)", level=3)
    doc.add_paragraph(
        "Review the Curry–Howard functional interpretation, its success in explaining proof terms, and the pressure points"
        " that motivate richer equality data: the erasure of ‘how’ a judgement is obtained, the tension between extensional"
        " and intensional equality (Martin-Löf), and the need to separate definitional computation from propositional evidence."
    )

    doc.add_heading("§1.2 Labelled Natural Deduction: labels as first-class citizens (≈4 pages)", level=3)
    doc.add_paragraph(
        "Present the LND paradigm as a harmonious coupling of a functional calculus on labels with a logical calculus on formulas"
        " (Gentzen/Prawitz; Gabbay’s LDS). Emphasize abstraction/discharge as the bridge between the two dimensions, and explain"
        " why labelled deduction naturally records the history of deduction steps."
    )

    doc.add_heading("§1.3 The equality judgement and rewrite reasons (≈6 pages) — from Ruy Ch.6", level=3)
    doc.add_paragraph(
        "Explain the judgement a =_s b : D, where the index s is a composition of definitional equalities (β, η, ζ, …) tracking"
        " the reason for identification. Develop the definitional vs propositional equality distinction and the ‘missing entity’"
        " problem: standard equality judgements forget the contextual information needed for intensional analysis. Situate the"
        " proposal as promoting rewrite reasons from indices to first-class mathematical objects (computational paths)."
    )

    doc.add_heading("§1.4 Normalisation as meaning: the philosophical programme (≈3 pages) — from Ruy Ch.5,8", level=3)
    doc.add_paragraph(
        "Summarize proof-theoretic semantics and the normalisation-centred view of meaning (Gentzen → Prawitz → Dummett → Martin-Löf"
        " with Tait’s intensional interpretations). Motivate the guiding principle: the rewrite/normalisation structure of equality"
        " proofs is not merely metatheory but carries mathematical meaning."
    )

    doc.add_heading("§1.5 From 35 rules to 77: where this book begins (≈4 pages)", level=3)
    doc.add_paragraph(
        "Bridge from the LND_EQ rewrite-reason calculus (≈35 rewrite rules, terminating and confluent via Knuth–Bendix/Newman) to the"
        " enriched computational-path rewrite system (≈77 rules) supporting congruence, transport, and higher-dimensional structure."
        " Explain how scaling the rewrite theory from reasons-as-indices to paths-as-objects forces new coherence rules and yields new"
        " higher cells."
    )

    doc.add_heading("§1.6 The architectural choice: Path vs Eq (≈3 pages)", level=3)
    doc.add_paragraph(
        "State upfront the two-level architecture: Eq for proof-irrelevant propositional equality, Path for proof-relevant traces."
        " Clarify the status of UIP: it may hold at the Eq-level in classical foundations but fails for Path (distinct traces)."
        " Compare with two-level type theories and explain what the refinement buys: explicit coherence witnesses, decidable equality"
        " of normal forms, and a rewriting-driven route to higher groupoids."
    )

    doc.add_heading("Chapter 2 — Labelled Natural Deduction and Proof-Terms", level=2)
    doc.add_paragraph(
        "A self-contained presentation of LND: proof terms for connectives/quantifiers, β/η/ζ reductions and the role of Knuth–Bendix"
        " completion (e.g. the discovery of additional reductions to restore confluence)."
    )

    doc.add_heading("Chapter 3 — Computational Paths: Definitions and Elementary Algebra", level=2)
    doc.add_paragraph(
        "Define computational paths as structured witnesses built from rewrite steps; introduce refl/trans/symm/congruence and their"
        " basic interactions. Establish the proof-relevance of Path and explain the intended quotient by rewrite equivalence."
    )

    doc.add_heading("Chapter 4 — The Rewrite System on Paths (≈77 rules)", level=2)
    doc.add_paragraph(
        "Present the rewrite system governing path expressions, organized by structural, functorial, congruence, and transport rules,"
        " and define the induced equivalence relation on paths."
    )

    doc.add_heading("Chapter 5 — Termination, Confluence, and Normal Forms", level=2)
    doc.add_paragraph(
        "Prove termination and confluence (Newman’s lemma; critical-pair analysis), yielding unique normal forms and a practical equality"
        " test for path expressions via normalisation."
    )

    # Part II
    doc.add_heading("Part II — Higher Structure: From Groupoids to ω-Groupoids", level=1)

    doc.add_heading("Chapter 6 — The Groupoid of Computational Paths", level=2)
    doc.add_paragraph(
        "Construct the fundamental groupoid: objects are terms, morphisms are paths modulo rewrite equivalence; prove groupoid laws"
        " as theorems of the rewrite system."
    )

    doc.add_heading("Chapter 7 — Two-Cells: Derivations Between Paths", level=2)
    doc.add_paragraph(
        "Define 2-cells as rewrite derivations between paths; show vertical/horizontal composition and the interchange law (a 2-categorical"
        " structure over endpoints)."
    )

    doc.add_heading("Chapter 8 — Three-Cells and Coherence (Contractibility at ≥3)", level=2)
    doc.add_paragraph(
        "Introduce 3-cells between 2-cells and establish the key coherence/contractibility phenomenon at dimensions ≥ 3, isolating precisely"
        " the non-trivial structure at dimension 2 needed for interesting fundamental groups."
    )

    doc.add_heading("Chapter 9 — The Weak ω-Groupoid Theorem", level=2)
    doc.add_paragraph(
        "Assemble the globular tower and verify the axioms of a weak ω-groupoid (Batanin/Leinster style). Compare to the classical identity-type"
        " ω-groupoid results of Lumsdaine and van den Berg–Garner."
    )

    doc.add_heading("Chapter 10 — Eckmann–Hilton with Explicit Coherence", level=2)
    doc.add_paragraph(
        "Carry out Eckmann–Hilton in the computational-path setting with explicit coherence witnesses derived from rewriting, establishing commutativity"
        " of the second loop space."
    )

    # Part III
    doc.add_heading("Part III — Connections to Type Theory and Homotopy", level=1)

    doc.add_heading("Chapter 11 — Path Induction (J) as a Theorem", level=2)
    doc.add_paragraph(
        "Derive the J-eliminator from contractibility of the based path space inside the computational-path framework; develop transport and dependent"
        " reasoning as derived constructions."
    )

    doc.add_heading("Chapter 12 — UIP, Proof-Relevance, and the Two-Level Picture", level=2)
    doc.add_paragraph(
        "Analyze precisely how UIP behaves: it may hold for Eq in classical foundations while failing for Path. Compare with Streicher’s K, Hedberg’s theorem,"
        " Hofmann–Streicher groupoid model, and two-level type theory perspectives."
    )

    doc.add_heading("Chapter 13 — Fundamental Groupoids of Standard Spaces", level=2)
    doc.add_paragraph(
        "Compute π₁ for circle, torus, Klein bottle, projective plane, bouquets; develop Seifert–van Kampen and selected covering space results in the path algebra."
    )

    doc.add_heading("Chapter 14 — Relationship to HoTT, Cubical, and Higher-Dimensional Rewriting", level=2)
    doc.add_paragraph(
        "Provide a careful dictionary and comparison: HoTT Book and Rijke; cubical type theory; higher-dimensional rewriting (Squier; Guiraud–Malbos). Emphasize that"
        " computational paths provide an algebraic route to familiar higher structures without adopting univalence as an axiom."
    )

    # Part IV
    doc.add_heading("Part IV — Higher Algebra and Applications", level=1)

    doc.add_heading("Chapter 15 — Categorical and Higher-Algebraic Perspectives", level=2)
    doc.add_paragraph(
        "Interpret computational paths as a free groupoid on a computad-like presentation; discuss enrichment, computads, and connections to ∞-categories and the homotopy"
        " hypothesis."
    )

    doc.add_heading("Chapter 16 — Perspectives and Open Problems", level=2)
    doc.add_paragraph(
        "Future directions: dependent extensions, directed paths (without symmetry), potential interaction with univalence, and the broader programme of extracting homotopy-theoretic"
        " invariants from rewriting systems."
    )

    # Appendices
    doc.add_heading("Appendices", level=1)

    doc.add_heading("Appendix A — Term Rewriting Systems (background)", level=2)
    doc.add_heading("Appendix B — Groupoids and Higher Groupoids (background)", level=2)
    doc.add_heading("Appendix C — Formal Verification Methodology (Lean 4)", level=2)
    doc.add_paragraph(
        "This appendix describes the computer-checked development as a methodological guarantee and guide to reproducibility. It is not required for understanding the mathematical"
        " development in the main text."
    )
    doc.add_heading("Appendix D — Complete Rule List (≈77 rules)", level=2)

    # Uniqueness / contributions
    doc.add_heading("What Makes This Book Unique (Five Contributions)", level=1)
    add_numbered(
        doc,
        [
            "A rewriting-driven derivation of groupoid and weak ω-groupoid structure from explicitly recorded equality reasons.",
            "A detailed, terminating and confluent rewrite system on path expressions (≈77 rules), with normal forms and a decidable equality test at the syntactic/path level.",
            "A proof that path induction (J) follows as a theorem from contractibility of the based path space, rather than being postulated.",
            "Explicit coherence data throughout (including Eckmann–Hilton), with named witnesses rather than existential appeals.",
            "A unified bridge between proof theory (LND), rewriting theory, and algebraic topology, with concrete computations of fundamental groups/groupoids.",
        ],
    )

    # Audience
    doc.add_heading("Target Audience", level=1)
    add_bullets(
        doc,
        [
            "Algebraists and topologists interested in algebraic models of homotopy types (groupoids, higher groupoids).",
            "Type theorists and proof theorists interested in equality, identity, and normalization.",
            "Researchers in term rewriting and higher-dimensional rewriting.",
            "Advanced graduate students with background in algebra and basic category theory; no proof-assistant knowledge required.",
        ],
    )

    # Literature
    doc.add_heading("Relationship to Existing Literature", level=1)
    doc.add_heading("Homotopy type theory and univalent foundations", level=2)
    doc.add_paragraph(
        "The HoTT Book and Rijke develop identity/path types within intensional/cubical settings. This proposal is complementary: it shows how comparable higher structures arise"
        " from rewrite traces attached to ordinary equality judgements. Univalence is not assumed; its interaction with computational paths is posed as an open problem."
    )
    doc.add_heading("Identity-type ω-groupoids", level=2)
    doc.add_paragraph(
        "We build on the paradigm that identity types carry weak ω-groupoid structure (Lumsdaine; van den Berg–Garner) and provide a distinct construction driven by explicit rewriting."
    )
    doc.add_heading("Rewriting theory and higher-dimensional rewriting", level=2)
    doc.add_paragraph(
        "We connect the concrete rewrite system on equality reasons to the rewriting–homotopy interface (Squier; Guiraud–Malbos) and provide a self-contained termination/confluence account."
    )

    # Publishers
    doc.add_heading("Recommended Publishers", level=1)
    doc.add_heading("First choice: Cambridge Tracts in Mathematics (Cambridge University Press)", level=2)
    doc.add_paragraph(
        "The book is a research monograph developing a coherent theory at the intersection of several mathematical areas; the Tracts series is a natural home for such work."
    )
    doc.add_heading("Other plausible venues", level=2)
    add_bullets(
        doc,
        [
            "Springer Lecture Notes in Mathematics (for a shorter version) or a Springer monograph series.",
            "Oxford Logic Guides (if positioned more strongly toward proof theory and type theory).",
        ],
    )

    # Selling points
    doc.add_heading("Summary of Selling Points", level=1)
    add_bullets(
        doc,
        [
            "A new, explicit mathematical route from labelled deduction and rewrite reasons to higher groupoids.",
            "Concrete, checkable rewrite-theoretic metatheory (termination, confluence, normal forms) underpinning the higher structure.",
            "Clarity about the two-level Path vs Eq architecture and how it relates to UIP and existing foundations.",
            "Nontrivial computations in algebraic topology within a rewriting-driven formalism.",
            "A large, coherent computer-checked development described as methodology, not as prerequisite.",
        ],
    )

    # Timeline
    doc.add_heading("Timeline", level=1)
    doc.add_paragraph(
        "Manuscript completion: 12–18 months from proposal acceptance. The mathematical core is developed; the principal work is exposition, organization, and polishing of the narrative"
        " for a broad mathematical audience."
    )

    doc.add_heading("Estimated Length", level=1)
    doc.add_paragraph("Target length for publication: 400–500 pages (including appendices).")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("out", nargs="?", default=str(DEFAULT_OUT), help="Output .docx path")
    ap.add_argument("--template", default=None, help="Optional .docx template to preserve publisher styles/headers")
    args = ap.parse_args()

    out = Path(args.out)
    template = Path(args.template) if args.template else None
    main(out, template)
    print(str(out))
