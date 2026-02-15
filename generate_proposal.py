#!/usr/bin/env python3
"""Generate a polished Word (.docx) book proposal from book_proposal_opus.md.

Creates:
  /Users/arthur/clawd/ComputationalPathsLean/book_proposal.docx

Design goals (per instructions):
- Professional Word formatting (Heading 1 for parts, Heading 2 for chapters, Heading 3 for sections)
- Include the full proposal content from book_proposal_opus.md (all 16 chapters + appendices)
- Expand Chapter 1 with substantial exposition (pure mathematics; Lean only in Appendix C)

Note: The Chapter 1 expansion is intentionally long, to approximate the requested 20–25 page depth.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path("/Users/arthur/clawd/ComputationalPathsLean")
IN_MD = ROOT / "book_proposal_opus.md"
OUT_DOCX = ROOT / "book_proposal.docx"

TITLE = "Computational Paths and the Algebraic Topology of Rewriting: From Labelled Deduction to Higher Groupoids"


# -----------------------------
# Word styling helpers
# -----------------------------

def set_default_styles(doc: Document) -> None:
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # Headings (keep Word defaults, but ensure consistent font)
    for sty in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        if sty in doc.styles:
            s = doc.styles[sty]
            s.font.name = "Times New Roman"

    # Page margins (approx standard)
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Book Proposal")
    r.font.size = Pt(13)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Authors: [to be confirmed — Ruy de Queiroz; Arthur …]")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Prepared: {date.today().strftime('%B %Y')}")

    doc.add_page_break()


# -----------------------------
# Markdown-ish parsing helpers
# -----------------------------

def strip_inline_md(s: str) -> str:
    # Convert markdown links [text](url) -> text (url)
    s = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1 (\2)", s)
    # Remove emphasis/code markers
    s = s.replace("**", "")
    s = s.replace("*", "")
    s = s.replace("`", "")
    # Normalize whitespace
    s = re.sub(r"\s+", " ", s).strip()
    return s


def add_paragraphs_from_block(doc: Document, block: str) -> None:
    block = block.strip("\n").rstrip()
    if not block.strip():
        return
    parts = re.split(r"\n\s*\n", block)
    for para in parts:
        t = strip_inline_md(" ".join(line.strip() for line in para.splitlines()).strip())
        if t:
            doc.add_paragraph(t)


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        it = strip_inline_md(it)
        if it:
            doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        it = strip_inline_md(it)
        if it:
            doc.add_paragraph(it, style="List Number")


# -----------------------------
# Chapter 1 expansion content
# -----------------------------

def add_ch1_expansion(doc: Document) -> None:
    """Insert the requested expanded Chapter 1 sections.

    This section is intentionally substantial: it is meant to be read as a *book-level*
    Chapter 1 expansion (≈20–25 pages in a typical manuscript layout), giving deeper
    coverage of de Queiroz’s labelled deduction programme and explaining how the
    computational-paths viewpoint emerges.

    It is written as pure mathematics/proof theory/higher algebra. Lean appears only
    later in Appendix C.
    """

    def h3(title: str) -> None:
        doc.add_heading(title, level=3)

    def block(text: str) -> None:
        add_paragraphs_from_block(doc, text)

    h3("§1.1 Curry–Howard correspondence (≈3pp)")
    block(
        """
        The Curry–Howard correspondence is the guiding dictionary that permits a single mathematical object to be read in two complementary ways: as a derivation in a proof system and as a term in a typed calculus. Propositions correspond to types; proofs correspond to programs; proof normalisation corresponds to computation; and proof equality corresponds to program equivalence.

        For the present book, the essential lesson is not merely that proofs can be represented by λ-terms, but that the *transformations* of proofs—β-reduction, η-expansion, commuting conversions, and the normalisation procedures they generate—carry structural information that is normally relegated to metatheory. When one decides to make that structural information explicit, the equality theory of proofs takes on an independent mathematical life.

        A concrete starting point is the simply typed λ-calculus and natural deduction for implicational logic. Implication introduction corresponds to abstraction (λx.t), and implication elimination corresponds to application (t u). The β-reduction (λx.t) u → t[x := u] is the computational content of eliminating a detour: an introduction immediately followed by an elimination. In categorical semantics, this corresponds to the universal property of exponentials.

        The same pattern repeats for other connectives. Conjunction corresponds to products: introduction is pairing ⟨t,u⟩; elimination is projection π₁, π₂; and β-reduction states π₁⟨t,u⟩ → t and π₂⟨t,u⟩ → u. Disjunction corresponds to coproducts: introduction is injection inl/inr; elimination is case analysis; and β-reductions express the computation of case on injections.

        These reduction rules do two jobs simultaneously.

        (1) They ensure canonicity and computational adequacy: closed proofs of a canonical proposition reduce to canonical forms.

        (2) They generate a notion of equality on proofs/programs: two terms are definitionally equal if they are connected by a chain of β/η conversions.

        Standard presentations treat definitional equality as “built into the judgement”: one may rewrite terms freely, without tracking the steps. This is practical, but it hides a key mathematical phenomenon: definitional equality is itself generated by *rewriting*, and rewriting has higher-dimensional structure.

        To see this, observe that a term may have multiple reduction paths to a normal form. Confluence states that all such paths ultimately agree in their endpoint (up to further reduction), but it does not say that the *paths* are identical. A choice is therefore made: either one collapses all reduction histories (proof-irrelevant equality), or one remembers them (proof-relevant equality). The second choice is what opens the door to the algebraic topology of rewriting.

        There are three specific pressures within Curry–Howard that force one to confront this choice.

        First: substitution is functorial and interacts with reduction. If one has a reduction t → t′, then substitution t[x := u] → t′[x := u] should hold. In other words, reductions must be stable under context and under term constructors. This stability is precisely what rewriting theorists call “compatibility” (closure under contexts).

        Second: extensionality principles appear as η-rules. For functions, η states that any f is extensionally equal to λx. f x. In a judgemental setting, η can be built into definitional equality, but again, one must decide whether to record this conversion explicitly.

        Third: equality becomes a *connective of its own*. In intensional type theory, equality is internalised as an identity type Id_A(a,b), and its elimination principle J encodes path induction. The higher groupoid structure on identity types (in the sense of Hofmann–Streicher, Lumsdaine, van den Berg–Garner) is a theorem about the rules of Id; but those rules are themselves axioms of the theory.

        In contrast, the computational-paths programme starts from an independently motivated proof-theoretic device: rewrite reasons in labelled deduction. The immediate output is a structured notion of equality witness, built from the same primitives (β, η, congruence) that already occur in Curry–Howard. The surprise is that, when treated as objects, these witnesses organise into groupoids and, by iterating rewriting, into weak ω-groupoids.

        One can summarise the transition as follows. Curry–Howard supplies syntax (typed terms as proofs). Labelled deduction supplies an equality judgement with *reasons* (rewrite traces). Rewriting theory supplies an analysis of those traces (termination, confluence, critical pairs). Higher algebra supplies the correct language to package the resulting coherence (groupoids and higher groupoids). The present book develops this transition carefully and proves that it is not merely an analogy: the algebra of rewriting yields genuine homotopical structures.

        Historically, this story connects several traditions.

        • From Gentzen and normalization (cut-elimination) through Prawitz (natural deduction normalization) to the proof-theoretic semantics programme.

        • From Martin-Löf’s intensional equality to the groupoid model of Hofmann–Streicher and the later homotopy-theoretic interpretation of identity types.

        • From abstract rewriting systems to higher-dimensional rewriting (Squier and subsequent work), where confluence data is understood homotopically.

        The thesis of this book is that the meeting point of these traditions is not accidental. Once equality reasons are made explicit, the groupoid operations are already present syntactically, and higher-dimensional coherence arises as a systematic consequence of rewriting.

        A minimal example already exhibits this phenomenon. Consider a term t and the composite (reduction) steps that witness that t is definitionally equal to itself: one may take the empty reduction (reflexivity), or perform an η-expansion followed by a β-reduction, or apply a commuting conversion twice, and so on. At the level of mere definitional equality, all such “proofs of equality” are collapsed: one says simply that t ≡ t. But if the equality witness is retained, then one immediately has a nontrivial space of loops based at t, generated by basic conversions and their composites.

        The moment one admits composition of equality witnesses, one confronts associativity issues. In a judgemental treatment, associativity is never mentioned: rewriting “modulo associativity” is done tacitly. If witnesses are objects, however, there are distinct composites corresponding to different parenthesisations. The fact that these composites should be identified is itself a mathematical statement, and it is expressed by a second-level rewrite (an associator), which is a 2-dimensional cell.

        In the typed λ-calculus, the “congruence closure” of reduction—the principle that one may reduce inside any context—creates further coherence requirements. If one reduces t → t′, then f t → f t′ for any term former f. The induced action on equality witnesses must preserve identity, composition, and inversion; i.e., it must be functorial. This functoriality is again something that is automatic only if one has decided not to regard witnesses as objects.

        These coherence considerations can be phrased categorically: proofs/terms form a category; equality witnesses suggest a groupoid enrichment; and the rules of the calculus supply the generators and relations of this enriched structure. They can also be phrased homotopically: terms are points, reductions are paths, and confluence diagrams are homotopies.

        The book’s contribution is to treat this structure systematically and with proofs. Rather than assuming “paths in a space” or postulating identity-type rules, we begin with the proof-theoretically motivated rewrite system implicit in Curry–Howard/LND and show that it generates, by iteration, precisely the kind of weak ω-groupoid structure expected from the homotopy hypothesis.
        """
    )

    h3("§1.2 Labelled Natural Deduction (≈4pp)")
    block(
        """
        Labelled Natural Deduction (LND) is a framework in which every judgement is paired with a label: a term that records the computational content of the derivation. A rule of inference therefore has two simultaneous components: a logical rule on formulas and a term-former on labels.

        The conceptual advantage of LND is that it separates and coordinates three layers that are often collapsed.

        • The logical layer: derivations Γ ⊢ A built from introduction and elimination rules.

        • The computational layer: proof terms a,b,… that label derivations and support a term calculus with substitution.

        • The equality layer: a structured account of when two labels/derivations are equal, with explicit reasons.

        In a standard natural deduction system, one may think of a derivation as an unlabelled tree. LND enriches the tree by propagating a label through the derivation. For implication, the introduction rule discharges an assumption and produces an abstraction on labels; the elimination rule applies one label to another. For conjunction, introduction builds a pair; elimination projects. For disjunction, introduction injects; elimination performs case analysis. For quantifiers, introduction and elimination interact with eigenvariable conditions and substitution.

        A central mathematical property of such systems is subject reduction: if a label a has type A and reduces to a′, then a′ also has type A (in the appropriate typing judgement). This ensures that reduction is sound with respect to the logical reading.

        De Queiroz’s programme refines this by focusing on the *equalities* induced by normalisation. In many proof-theoretic accounts, normalisation is a relation on derivations: two derivations are “the same proof” if they normalise to the same normal form. LND makes visible that the normalisation procedure is reflected in the term language by β/η and related conversions. Thus one can speak of equality of proofs by speaking of equality of their labels.

        The critical innovation is the explicit equality judgement indexed by a reason. Instead of writing merely a = b : A, one writes

            a =_s b : A,

        where s is built from elementary conversion steps (β, η, commuting conversions, and structural substitutions). The role of s is to *record* the computational path from a to b.

        This record is not a mere list: it is closed under operations. Reasons can be composed (transitivity), inverted (symmetry), and transported through term constructors (congruence). Moreover, to preserve a meaningful equality theory, the calculus of reasons must itself satisfy equations: associativity of composition, unit laws, inverse laws, and functoriality of congruence.

        At a practical level, one may think of s as a term in a second calculus—an “equality term calculus”—whose constructors correspond to the logical principles governing equality: reflexivity, symmetry, transitivity, congruence, and (in dependent contexts) transport.

        At a mathematical level, the calculus of reasons has exactly the shape of a path algebra.

        • Objects: terms/proof labels.

        • 1-cells: reasons s witnessing identifications.

        • Composition: transitivity.

        • Inverses: symmetry.

        When the associated rewrite rules on reasons are added (eliminating bureaucratic parentheses and cancelling inverses), one recognises the axioms of a groupoid up to a further definitional equality.

        LND therefore provides the syntactic birthplace for computational paths: it is the environment in which equality reasons arise naturally and are operationally meaningful.

        The present book takes LND as a starting point, but it aims at a different endpoint. The equalities in LND were originally designed to account for proof identity and normalisation. Here we show that the same formalism, when taken seriously as an object of algebraic study, yields the structures that homotopy theorists identify in spaces: fundamental groupoids, higher homotopies, and weak ω-groupoids.

        A reader familiar with category theory may use the following analogy. A proof system is like a presentation of a free categorical structure. Proof terms correspond to morphisms; rewriting corresponds to identifying morphisms by specified relations; confluence corresponds to coherence of those identifications. LND makes the generating morphisms and relations explicit and therefore allows one to compute the resulting coherence data.

        This is not an abstract slogan. In later chapters, the rewrite system on computational paths is analysed by standard rewriting-theoretic techniques: well-founded measures for termination, critical pair analysis for local confluence, and Newman's lemma for confluence. The entire higher structure is then extracted from this rewrite analysis.

        To make the LND picture more concrete, one may view a labelled derivation as a judgement of the form Γ ⊢ a : A, where A is the formula/type proved and a is the label encoding the computational content. The derivation rules then specify simultaneously (i) how A is formed and (ii) how a is built. Equality is tracked by a judgement Γ ⊢ a =_s b : A, where s is a term of the reason calculus.

        The reason calculus is designed to mirror the “obvious” operations on equality proofs that practitioners already use informally. One composes equalities (to substitute equals for equals), reverses them (symmetry), and applies them under constructors (congruence). The innovation is not the existence of these operations but their explicit, typed representation.

        Once reasons are explicit, two further questions become unavoidable.

        (1) Which equalities between reasons should be admitted? In other words, when do two reasons s and s′ count as “the same” reason? This is where the rewrite system on reasons is introduced.

        (2) What is the correct balance between proof-relevance and proof-irrelevance? If one identifies too much, one collapses interesting structure; if one identifies too little, one is left with an unwieldy syntax. The book’s answer is to impose precisely the coherence rewrites needed to obtain categorical laws (groupoid laws, interchange laws) while retaining nontrivial low-dimensional information.

        It is this balance that makes the proposal relevant to mathematicians outside logic. The point is not to advocate a new proof formalism; it is to extract an algebraic object (a groupoid and higher groupoid) from a familiar proof-theoretic setting and to study it with the tools of algebra and topology.

        In particular, LND supplies a disciplined mechanism for “keeping track of the reason”. Once that bookkeeping is regarded as an algebraic structure, one can compare it with the fundamental groupoid of a space, compute loop groups, and apply familiar topological arguments (such as Eckmann–Hilton) in a syntactic setting.

        From the standpoint of exposition, Chapter 2 of the book will present the LND calculus in a self-contained manner: the term language for each connective, the typing rules, and the associated computational conversions (β and η families). Particular emphasis is placed on the role of discharge/abstraction, because it is there that the proof-theoretic structure controls substitution and therefore controls the stability properties required for the later path algebra.

        For example, implication introduction corresponds to taking a derivation of B from an additional assumption A and turning it into a derivation of A → B; on labels, this is λ-abstraction. Equality reasons must respect this operation: if two labels are equal under a reason s, then their abstractions are equal under a transported reason. Similar remarks apply to pairing and projections, injections and case analysis, and quantifier rules.

        LND also provides a natural setting for discussing the “geography” of equality: which equalities are definitional (computational), which are propositional (logical), and which are meta-level equivalences between derivations. The reason-indexed equality judgement can be read as a bridge between these notions: it carries computational content but is expressed as a judgement whose proof-theoretic behaviour can be studied.

        This bridge is particularly important for mathematicians who come to the subject from topology or algebra rather than from type theory. For such readers, LND is not primarily a foundational system; it is a presentation that makes a rewriting algebra visible. Once visible, that algebra can be analysed with familiar tools: presentations by generators and relations, quotient constructions, and the extraction of invariants.

        It is also worth stressing what LND *does not* require. One does not need to adopt a particular foundational stance on the identity type or on univalence in order to work in LND. The system can be presented in a classical metatheory and treated as a combinatorial object: a typed term system together with a labelled equality judgement.

        This makes LND especially suitable for a “Cambridge Tracts” style audience. A reader may accept the term calculus and rewrite rules as definitions, just as one accepts a presentation of a group or a monoidal category by generators and relations, and then proceed to study the resulting algebra.

        Finally, LND clarifies the role of *context* (Γ). In proof theory, context controls assumptions and discharge. In rewriting, context controls where reductions may be performed (closure under contexts). In higher algebra, context corresponds to functoriality: operations must respect composition. LND is the point at which these three notions align. This alignment is one of the reasons the computational-path groupoid and higher groupoid constructions are not artificial: they are the natural algebraic completion of the stability requirements already implicit in labelled deduction.

        Concretely, “closure under contexts” is the statement that if a rewrite step transforms a into b, then placing that step inside any surrounding term context yields a rewrite step of the larger term. In categorical language, this is the statement that term constructors act functorially. In proof-theoretic language, this is the stability of proof transformations under substitution and under the formation of compound deductions.

        Once one records reasons explicitly, this stability becomes an algebraic structure: one must define how a reason is mapped by each constructor and prove that the mapping respects composition and inversion. These are precisely the functoriality laws that later show up as congruence rules in the path rewrite system.

        In this way, LND provides more than motivation: it provides the blueprint for the rule families that must exist in any coherent computational-path calculus. The “extra” rules beyond the base βη-equality rules are not arbitrary additions; they are forced by the requirement that equality reasons behave well under the operations that create and manipulate proofs.
        """
    )

    h3("§1.3 The equality judgement and rewrite reasons (≈6pp)")
    block(
        """
        The equality judgement a =_s b : D is the conceptual hinge of the book. It replaces the bare assertion “a equals b” by a triple (a,b,s) in which s is an explicit witness—a structured reason recording how the identification is obtained.

        One may view the move from “a = b” to “a =_s b” as analogous to the move from a mere reachability relation in a graph to an explicit path object. In a graph, reachability is a proposition; a path is data. Rewriting theory and homotopy theory are both fundamentally about what happens when one studies path data instead of mere reachability.

        The elementary generators of reasons come from the computational rules already present in the proof-term calculus: β-reductions, η-expansions, commuting conversions, and structural conversions for substitution. A reason is then built by composing these generators.

        De Queiroz introduces a small algebra of reason constructors (a schematic presentation is):

        • ρ: reflexivity, witnessing a = a.

        • σ(s): symmetry, reversing s.

        • τ(s,t): transitivity, concatenating s and t.

        • congruence/whiskering operations, carrying reasons through term formers.

        This list already suggests a groupoid: ρ is an identity, σ provides inverses, and τ provides composition. But at the raw syntactic level, groupoid laws do not hold strictly. For instance, τ(τ(s,t),u) and τ(s,τ(t,u)) are distinct expressions. The groupoid laws therefore appear as *rewrites between reasons*.

        Here the distinction between two equalities becomes crucial.

        • There is an equality relation on labels, witnessed by reasons (a =_s b).

        • There is an equality relation on reasons themselves (s ≡ s′), generated by rewrite rules that express coherence.

        The second relation is where the higher structure lives. Associativity, unit laws, and inverse cancellation become rewrite rules that reduce reason expressions to normal forms. In other words: the groupoid laws are not axioms; they are *theorems of a rewrite system*.

        This is one of the central ideas of the book: coherence is implemented as rewriting.

        To make this concrete, consider the following typical equations (schematically).

        • Unit laws: τ(ρ, s) → s and τ(s, ρ) → s.

        • Inverse laws: τ(σ(s), s) → ρ and τ(s, σ(s)) → ρ.

        • Associativity: τ(τ(s,t),u) → τ(s,τ(t,u)) (or the reverse orientation, depending on the chosen normal form discipline).

        • Involution: σ(σ(s)) → s.

        These are precisely the reductions one expects when manipulating paths in a fundamental groupoid. But here they arise as rewrites on equality reasons.

        Once one recognises reasons as 1-dimensional paths, it is natural to ask for 2-dimensional structure: what is a witness that two reasons s and s′ are equal? In rewriting-theoretic terms, such a witness is a *derivation*: a sequence of rewrite steps transforming s into s′. In the language of the proposal, such derivations are 2-cells.

        This iteration produces an entire tower.

        • 1-cells: computational paths (reasons).

        • 2-cells: derivations showing that two paths are rewrite-equivalent.

        • 3-cells: derivations between derivations (coherence between coherence proofs).

        • and so on.

        A potential worry is that this tower might collapse into triviality: if we freely identify too much, do we destroy the interesting information (such as nontrivial fundamental groups)? The book addresses this explicitly by isolating which identifications are imposed at which level.

        The quotient on 1-cells is chosen to remove the “bureaucracy” of parenthesisation and immediate cancellations while preserving computational content. The resulting quotient is well-behaved enough to satisfy groupoid laws, yet fine-grained enough to distinguish genuinely different rewrite traces when appropriate.

        The key technical move is to work with a terminating and confluent rewrite system on path expressions. Termination ensures that every path expression reduces to some normal form; confluence ensures that the normal form is unique. Together, they yield a canonical representative of each equivalence class, which makes equality decidable at the syntactic/path level.

        This is the point at which rewriting theory and algebraic topology meet. In topology, one studies paths up to homotopy; homotopies themselves can be compared up to higher homotopy; and the resulting structure is a fundamental ω-groupoid. In rewriting, one studies reduction sequences up to a chosen congruence; confluence diagrams are 2-cells; and higher coherence expresses contractibility of spaces of reductions. The computational-paths framework makes this parallel precise.

        To connect to the later chapters, we stress two features.

        (1) Non-UIP at the path level. If paths are rewrite traces, then different traces between the same endpoints can exist and need not be identified. This yields a proof-relevant equality object. The book later exploits this to obtain nontrivial fundamental groups.

        (2) Explicit coherence. Rather than postulating associators, unitors, and interchangers as axioms, the book constructs them as explicit derivations arising from the rewrite system. This makes the resulting higher groupoid “computationally explicit” in a rewriting-theoretic sense.

        The conclusion of this section is therefore a methodological one: by moving from equality-as-proposition to equality-with-reason, and by allowing rewriting at the level of reasons, one automatically obtains the algebraic infrastructure required for higher groupoids.

        It is worth emphasizing that the move is not merely terminological. It changes what it means to prove an equality statement. Under a reason-indexed equality judgement, proving a = b requires producing not only the endpoint identification but also the *route* by which the identification is achieved.

        Consider a simple rewriting situation familiar from λ-calculus. Let t be a term with two redexes, one nested in the other, or one in each component of a pair. One may reduce the left redex first, then the right; or the right first, then the left. Confluence guarantees that the two strategies eventually meet, but the *diagram* witnessing their meeting is a 2-dimensional piece of data. In rewriting-theoretic terms it is a join of a peak; in homotopical terms it is a homotopy between two paths.

        In the computational-path setting, this diagram becomes a 2-cell between two 1-cells (two reasons). The 2-cell can itself be rewritten or composed with other 2-cells (vertical composition corresponds to concatenating derivations; horizontal composition corresponds to whiskering). At this point the first genuinely higher-categorical structure appears: an interchange law relating vertical and horizontal composition.

        The book makes this precise by defining, for each pair of terms (a,b), a type/object of computational paths Path(a,b), and then defining a relation RwEq on paths generated by the rewrite rules on path expressions. A “derivation” between paths p and q is precisely a witness that p and q are related by RwEq; such derivations become the 2-cells of the theory.

        The reader should notice two important design principles.

        (A) Separating generators from relations. The path constructors (refl, symm, trans, congruence) generate a free syntax of paths. The rewrite rules then impose relations that express coherence. This is analogous to presenting a group by generators and relations: without a presentation one cannot compute.

        (B) Orienting relations as reductions. Rather than merely stating equations, we orient them into a rewrite system and prove it terminating and confluent. This produces canonical normal forms and, therefore, a decision procedure for equivalence of path expressions.

        The choice to orient coherence equations is mathematically consequential. It ensures that computations in the path algebra are tractable and that the resulting quotient structure is well-defined. It also aligns the development with the standard rewriting literature (critical pairs, completion, normalization theorems).

        We also stress that the groupoid structure is not assumed. It is derived. For example, the groupoid law (p · q) · r = p · (q · r) is proved by exhibiting a rewrite derivation from trans(trans(p,q),r) to trans(p,trans(q,r)). Similarly, the unit and inverse laws are proved by reducing trans(refl,p) and trans(p,refl) to p, and trans(symm(p),p) to refl, using the rewrite rules.

        Once these derivations are explicit, they can be reused as higher coherence data. The associator is not just a proposition that associativity holds; it is a concrete 2-cell. The same holds for unitors and inverse laws. In later chapters, these become the building blocks of the ω-groupoid structure.

        Finally, we note how this framework relates to the standard identity-type story. In intensional type theory, one postulates a type Id_A(a,b) of equalities and derives a groupoid structure on it. Here, we start with a syntactic equality witness (a rewrite reason) arising from proof normalisation and define the corresponding Path(a,b). The higher structure then emerges by iterating rewriting. Thus computational paths provide an *algebraic* route to the same kind of higher groupoid phenomena, with explicit normalisation and confluence theorems grounding the construction.

        A worked schematic example may help. Suppose we have a product type A × B with projections π₁, π₂ and pairing ⟨_,_⟩. The β-rules state π₁⟨x,y⟩ → x and π₂⟨x,y⟩ → y. Now fix terms u : A and v : B and consider the term π₁⟨u,v⟩. There are at least two “ways” in which this term is equal to u.

        • One may reduce directly by the β-rule for π₁, producing u.

        • One may first η-expand ⟨u,v⟩ in a context (depending on the chosen η-principles), then perform β-reductions, then simplify.

        Both routes establish definitional equality, but they yield distinct equality reasons. If one retains reasons as objects, these routes become distinct paths. The rewrite system on reasons then provides a way to relate these paths: a derivation that transforms one reason expression into another is a 2-cell.

        In more complex examples, such as nested pairs or iterated applications, the space of reasons can be large. Confluence guarantees that different reduction strategies lead to a common result, but confluence itself is witnessed by a family of 2-cells (one for each critical pair). The computational-paths programme treats these 2-cells as mathematical objects that can be composed, compared, and transported.

        This viewpoint clarifies the role of “equations between equations”. In ordinary mathematics one writes an equation and then, if necessary, shows that it is compatible with operations (e.g. addition, multiplication, substitution). In the present setting, one must also show that *reasons* are compatible with operations. If p : a = b is a path and f is a function symbol/constructor, then there is an induced path ap_f(p) : f(a) = f(b). The equations ap_f(p · q) = ap_f(p) · ap_f(q) and ap_f(p^{-1}) = (ap_f(p))^{-1} are not tautologies at the syntactic level; they are coherence statements witnessed by rewrites.

        Such coherences are precisely what is required to interpret Path as the hom-set of a groupoid (and Derivation₂ as the hom-set of a 2-groupoid-like structure). When these coherences are made explicit, familiar topological arguments become available. For example, one can define a loop space Ω(A,a) as the set of paths a → a modulo the chosen equivalence, and one can define a second loop space Ω²(A,a) as 2-cells from refl to refl. The Eckmann–Hilton argument then becomes an argument about interchange and units in this 2-dimensional algebra.

        Another important point is that this framework is compatible with standard algebraic-topological intuition. A fundamental groupoid is built from paths modulo homotopy; here, one builds an analogous structure from rewrite traces modulo rewrite equivalence. The “homotopies” are derivations between traces. The distinction between strict and weak laws corresponds to the distinction between definitional equality (syntactic equality) and equality witnessed by a higher cell (rewrite derivation).

        In summary, the reason-indexed equality judgement does two things simultaneously: it refines equality to a proof-relevant object, and it provides the raw material from which higher coherence is obtained by rewriting. The rest of the book develops the necessary rewrite-theoretic metatheory and packages the resulting tower into a weak ω-groupoid with applications in algebraic topology.

        Two forward pointers connect this section to the rest of the proposal.

        First: the derived J-eliminator. A standard way to obtain path induction in homotopy type theory is to prove that the based path space at a point a is contractible: the type Σ(y : A), Path(a,y) has a canonical center (a, refl) and every other element is connected to it. In the computational-path framework, one can mimic this argument at the Path-level. The “contractibility proof” is itself built from explicit path constructions and rewrite derivations, and from it one extracts a J-like eliminator as a theorem. The significance is conceptual: path induction is not postulated, but derived from the algebra of rewriting.

        Second: the low-dimensional/nontrivial vs high-dimensional/contractible pattern. The computational-path ω-groupoid is not intended to be “wild” in all dimensions. Instead, the theory is engineered so that higher coherence (from dimension 3 upward) becomes contractible in the appropriate sense, while dimension 2 retains enough structure to encode nontrivial π₁ information. This is what allows classical topological invariants to appear while keeping the higher coherence manageable.

        These points are important for positioning. They show that computational paths are not a mere repackaging of identity types. They are a rewriting-driven construction that yields, with explicit normalisation and confluence theorems, the same kind of higher groupoid structures that homotopy theorists expect, together with a clear explanation of where nontrivial invariants live.

        As a further guide for the reader, one may keep the following “dimension table” in mind.

        • Dimension 0: terms/proof labels (objects).

        • Dimension 1: computational paths (rewrite traces) between terms.

        • Dimension 2: derivations witnessing rewrite equivalence between paths (confluence/critical-pair data; whiskering and interchange).

        • Dimension ≥3: coherences between derivations; these become contractible in the weak ω-groupoid sense.

        This table explains both why the theory is homotopical and why it remains computable. The nontrivial structure is concentrated where topologists expect it (π₁ and related low-dimensional invariants), while higher coherence is tamed by contractibility.

        From the standpoint of a Cambridge Tracts-style exposition, this allows the book to present explicit low-dimensional calculations (circle, torus, Klein bottle, projective plane, bouquets) while keeping the higher-dimensional categorical overhead under control. The reader sees the algebraic topology emerge from rewriting, rather than being asked to accept an abstract ∞-categorical framework at the outset.

        The chapters that follow therefore proceed in the natural mathematical order: define the path expressions and rewrite system; prove termination/confluence; derive groupoid laws; define 2-cells and prove interchange; assemble the ω-groupoid theorem; and then apply the resulting machinery to classical computations.
        """
    )

    h3("§1.4 Normalisation as meaning (≈3pp)")
    block(
        """
        Normalisation is the proof-theoretic mechanism that explains why elimination rules are justified by introduction rules. In Gentzen’s sequent calculus this appears as cut-elimination; in natural deduction it appears as the removal of introduction–elimination detours. Prawitz’s normalisation theorem and the broader tradition of proof-theoretic semantics insist that these normal forms are not merely a technical convenience: they represent the canonical content of a proof.

        In the Curry–Howard reading, normalisation is computation. A proof term reduces to a normal form; the normal form represents the observable behaviour of the program. When two terms reduce to the same normal form, they are extensionally the same program. This is the standard justification for working modulo βη-equivalence.

        De Queiroz’s programme adds a crucial intensional refinement: the *reason* why two terms are equal matters, and normalisation gives a privileged source of such reasons.

        From the standpoint of rewriting theory, normalisation is a property of a rewrite system: termination ensures that every term reduces to a normal form; confluence ensures that the normal form is unique. When both hold, the rewrite system presents a deterministic computational semantics even if the reduction relation itself is nondeterministic.

        The equality-reason calculus is therefore built from two interacting rewrite systems:

        • The base reduction system on proof terms (β/η/commuting conversions).

        • The induced reduction system on equality reasons, encoding the algebra of reflexivity/symmetry/transitivity/congruence.

        The second system is necessary because “normalisation as meaning” is not stable under composition unless one accounts for coherence. Suppose one has three terms a,b,c connected by reasons s : a = b and t : b = c. Then τ(s,t) is a reason a = c, but the manner in which τ associates is not unique: τ(τ(s,t),u) and τ(s,τ(t,u)) represent the same abstract composite. If one insists that equality reasons are meaningful, one must also explain why these different composites should be regarded as the same reason (up to a canonical rewrite).

        This is precisely where the algebraic-topological picture is unavoidable. The space of reductions from a to b is like a path space. Confluence is the existence of fillers for certain squares (2-dimensional cells). Higher coherence is the existence of fillers for higher-dimensional boundaries. In other words, normalisation has a geometry.

        This viewpoint is familiar in higher-dimensional rewriting: Squier showed that convergent rewriting systems can be used to compute homotopical invariants, and later work develops polygraphic and computad-like presentations of higher categories. The computational-paths programme identifies a particularly natural convergent system (arising from equality reasons in labelled deduction) and shows that its higher-dimensional content aligns with ω-groupoid structures.

        Philosophically, this can be read as an intensional refinement of proof-theoretic semantics: meaning is not only the normal form but also the structured space of reasons by which normalisation proceeds.

        Mathematically, the consequences are concrete. Termination and confluence yield unique normal forms for path expressions, giving a decision procedure for path equality. Coherence derivations become explicit 2- and 3-cells, and later chapters show how these cells assemble into weak ω-groupoids.

        From the rewriting standpoint, the key lesson is that confluence proofs are themselves structured. To prove confluence one often proves local confluence (joinability of critical pairs) and combines it with termination via Newman's lemma. Each joinability proof is a square of reductions; in a higher-dimensional rewriting language, these squares are generating 2-cells of a polygraph. The collection of these squares, together with their higher relations, presents a homotopy type.

        In the computational-path system, this rewriting geometry is not imported from topology; it is produced internally by the algebra of equality reasons. The coherence data required to form groupoids and higher groupoids is therefore not an extra layer. It is already present in the confluence structure of the rewrite system.

        This perspective explains why the book repeatedly returns to termination/confluence results. They are not mere technicalities. They are what allow us to treat path expressions as a computational algebra with canonical representatives, and they are what produce explicit higher cells that can be composed and compared.

        In brief: normalisation provides the semantics; rewriting provides the combinatorics; confluence provides the 2-dimensional geometry; and the iterative closure of these constructions provides the higher-dimensional algebra.

        The reader may interpret this as a concrete instance of a general principle in higher algebra: whenever one has a presentation of a structure by generators and relations, one should expect coherence questions about the relations. In rewriting language, these coherence questions appear as overlaps of rewrite rules and the requirement that overlaps be joinable. In homotopy language, they appear as higher cells filling boundaries.

        In a terminating, confluent system, the “space of reductions” between two objects has a particularly tractable shape: every reduction path can be compared to the unique normal-form reduction, and critical pairs generate the essential 2-dimensional data. This is one reason convergent rewriting systems are so well suited to homotopical extraction.

        For computational paths, this tractability is essential. It allows us to define a quotient of paths by rewrite equivalence that is mathematically honest (well-defined) and computationally usable (decidable equality via normal forms). It also allows us to *name* the coherence witnesses that later serve as associators, unitors, and interchangers.

        In classical algebraic topology, one often computes π₁ by choosing a set of generators and relations (for example via van Kampen) and then analysing the resulting group presentation. Here the analogy is close: the rewrite system supplies generators (primitive rewrites) and relations (coherence rewrites), and the resulting path groupoid can be studied as a presented algebraic object. The later chapters exploit this to compute fundamental groups of standard spaces presented combinatorially.

        Thus normalisation is not only a semantic doctrine. It is the mechanism that turns rewriting into a source of computable homotopical invariants.
        """
    )

    h3("§1.5 From 35 rules to 77 (≈4pp)")
    block(
        """
        The phrase “from 35 rules to 77” is best understood as a rewriting-theoretic completion phenomenon.

        In the labelled deduction tradition, one begins with a set of elementary conversions corresponding to the computational behaviour of connectives. These conversions account for β-reduction (detour elimination), η-expansion (extensionality/canonicity), and certain commuting conversions (reordering of eliminations). A system of this kind can often be organised as a terminating, confluent rewrite system; if it is not initially confluent, one applies completion techniques (Knuth–Bendix style) to add the missing conversions.

        That base system supports an equality-reason judgement, but the reasons are still treated as indices. Once we promote reasons to objects (computational paths), we must also account for the algebra of operations on reasons.

        The enlarged system includes, in addition to the base computational conversions, a family of rules governing:

        (A) The structural algebra of paths: cancellation of symmetries against transitivity, elimination of redundant reflexivities, and a chosen associativity normalisation.

        (B) Congruence operations: functorial action of term constructors on paths (ap/congrArg), and the way this action interacts with composition and inversion.

        (C) Transport-like operations: stability of paths under contexts (“terms with a hole”), and naturality conditions ensuring that rewriting commutes with plugging into contexts.

        (D) Higher coherence generators: rules required to join critical pairs created by the interaction of (A)–(C) with the base conversions.

        The increase in the number of rules is therefore not a mere proliferation. It reflects the fact that a path algebra must be closed under the operations that define its intended categorical structure.

        From the perspective of algebra, one may say: the raw syntax freely generates a structure (a path magma with symmetry and congruence). The rewrite rules then impose relations that force this structure to satisfy the axioms of a groupoid (and later, a 2-category-like structure on 2-cells). The complete rewrite system is therefore a presentation of the intended algebraic object.

        The book treats this rewrite system as a central mathematical object. It is presented explicitly and analysed by standard rewriting tools.

        • Termination is established by a well-founded complexity measure on path expressions. Typically, the measure counts constructors and assigns weights so that every oriented rewrite strictly decreases the measure.

        • Local confluence is established by critical pair analysis: whenever two rewrite rules can overlap, one checks that the resulting peak can be joined.

        • Confluence then follows by Newman's lemma (for terminating systems).

        Once confluence is proved, every path expression has a unique normal form. This yields a practical and conceptual payoff: equality of paths modulo the rewrite theory becomes decidable by reduction to normal form.

        Crucially, the rewrite theory is calibrated so that it does not collapse all paths. It collapses only the “structural bureaucracy” (parenthesisation and immediate cancellations) required to make composition associative and unital up to the chosen equivalence. It leaves intact the possibility of distinct normal forms that represent genuinely different rewrite traces. This is why nontrivial fundamental groups can arise.

        In later chapters, this rewrite system is the engine that produces higher structure. The associator, unitors, inverse laws, and interchange properties are not postulated: they are witnessed by explicit rewrite derivations, constructed from the rules of the system.

        From an expository standpoint, the book organises the 77 rules by “type”:

        • Structural/groupoid rules (refl/trans/symm interaction; associativity normalisation; cancellation laws).

        • Functoriality rules (how congruence respects composition and inversion).

        • Congruence/naturality rules (how equalities behave under term constructors and contexts).

        • Transport-style rules (stability of path operations under substitution and dependent-like contexts).

        This classification is important for readers: it makes clear which rules correspond to familiar algebraic laws and which are completion/coherence rules required to maintain confluence.

        The larger rewrite system is also what permits a clean separation of concerns in the later chapters. Chapter 4 presents the rules and the induced equivalence relation. Chapter 5 proves the meta-theorems (termination and confluence). Only then do we pass to structural theorems: the groupoid theorem (Chapter 6) and the higher-cell tower (Chapters 7–9).

        In particular, confluence is not assumed when we build higher structure; it is proved once, and then reused. This is precisely the kind of metatheory expected in a pure mathematics monograph: one proves that the defining equations of a structure yield a consistent, well-behaved quotient, and then one studies the quotient.

        It is also important that the rewrite rules are oriented with a specific normal-form discipline in mind. For example, associativity of composition may be oriented so that compositions are always right-associated, turning any composite trans(trans(p,q),r) into trans(p,trans(q,r)). Unit laws may be oriented to erase explicit refl constructors where they are redundant, and inverse laws may be oriented to cancel immediate symm–trans patterns. These choices do not change the mathematical content (the induced equivalence relation is the same), but they make termination proofs and normal-form computations feasible.

        The move from a smaller rule set (e.g. one designed only to capture βη-equality of proof terms) to the larger rule set required for computational paths is therefore analogous to passing from a presentation of a monoid to a presentation of a groupoid with an explicit functorial action. One must add not only the inverse laws but also the coherence laws ensuring that inverses and congruence behave well with respect to composition.

        In rewriting terms, each new constructor on paths introduces new critical overlaps, and confluence forces corresponding joining reductions. The completed system is thus best viewed as the minimal coherent closure of the original equality-reason calculus under the operations we wish to perform on paths.

        The book’s Appendix D lists the full rule set and provides a road map for readers who want to see how the critical pairs arise. This is particularly valuable for referees: it makes the proof-theoretic machinery transparent and checkable, and it demonstrates that the higher coherence claims are grounded in a concrete, finite rewriting object.

        In addition, the rewrite presentation provides a practical benefit for readers who wish to experiment. Because normal forms are unique, one can perform explicit calculations in the path algebra by reduction. This is analogous to reducing a word in a presented group to a normal form (when a confluent rewriting system is available), and it makes the later topological computations feel concrete rather than abstract.
        """
    )

    h3("§1.6 The architectural choice: Path vs Eq (≈3pp)")
    block(
        """
        A recurring difficulty in the foundations of equality is that the word “equality” is used for two different mathematical notions.

        • Equality as a proposition: a = b is a truth-valued statement. Proofs of it, in classical mathematics, are normally taken to be proof-irrelevant.

        • Equality as a structured witness: a is identified with b by a specific computation, rewrite trace, or homotopy. Such witnesses are data, and different witnesses can carry different information.

        The computational-paths programme makes a deliberate two-level architectural choice designed to keep both notions available and to be explicit about which one is used where.

        Eq-level (propositional equality) is the ambient equality used to state ordinary theorems and to form quotients. It is compatible with classical practice: one may identify objects by equivalence relations without tracking witnesses. This is the level at which we speak of the quotient of path expressions by rewrite equivalence, and it is the level at which the book is written as ordinary mathematics.

        Path-level (computational paths) refines equality by equipping it with explicit witnesses: a path from a to b is a concrete object built from rewrite steps and constructors (reflexivity, symmetry, transitivity, congruence, etc.). Different paths between the same endpoints are not automatically identified.

        This distinction clarifies the status of UIP (Uniqueness of Identity Proofs). UIP asserts that, for any a,b, any two proofs of a = b are equal. In a setting that already treats equality proofs as proof-irrelevant, UIP holds trivially—but only because the theory has chosen to ignore equality data. In computational paths, UIP fails at the Path-level: different rewrite traces remain different objects. This failure is not a defect; it is precisely the mechanism that allows nontrivial homotopy groups to appear.

        The Path vs Eq architecture also clarifies what kind of “homotopy theory” is being extracted.

        • The 1-dimensional information (loops and their compositions) lives at the Path-level.

        • The 2-dimensional information (derivations between paths) lives at the level of rewrites between path expressions.

        • Higher-dimensional information (coherence between derivations) arises from iterating this construction.

        The book shows that, starting at dimension 3, the relevant higher cells become contractible in a precise Batanin/Leinster sense, yielding a weak ω-groupoid whose nontrivial information is concentrated in low dimensions—exactly as one expects in many algebraic-topological situations.

        Readers familiar with homotopy type theory may compare this to two-level and observational type theories, where one distinguishes a strict equality from a path-like equality. The present book, however, keeps the exposition in the classical language of rewriting, groupoids, and higher algebra. The two-level architecture is presented as a mathematical design principle: keep Eq as the ambient equality of mathematics, and introduce Path as a refined equality object when one needs to study the algebraic topology of computation.

        This design choice has two concrete payoffs.

        • It makes the coherence data explicit and computable: associators, unitors, and interchangers appear as rewrite derivations, not as abstract existential statements.

        • It provides a clear conceptual account of how structures akin to identity-type ω-groupoids can arise from proof theory without adopting the identity type rules as axioms.

        One may also view the two-level architecture as a response to a common referee question: “Where, exactly, is the homotopy theory?” The answer is that the homotopy theory is extracted from Path-level data (rewrite traces) and its higher rewrite structure, while the surrounding exposition uses Eq-level equality in the ordinary sense. This ensures that the book can be read as mainstream mathematics, with no foundational commitments beyond those required to talk about rewriting systems and quotients.

        The architecture also cleanly separates two kinds of statements.

        • Statements about computation: e.g. that a given rewrite system is terminating and confluent, that normal forms exist and are unique, or that a particular path expression reduces to another.

        • Statements about resulting algebraic objects: e.g. that the quotient of paths forms a groupoid, that 2-cells compose and satisfy interchange, or that a based path space is contractible (leading to a derived J-eliminator).

        This separation is important for the book’s positioning relative to HoTT. In HoTT, the identity type is primitive and its eliminator is axiomatic; higher groupoid structure is then derived. Here, the primitive data is a proof-theoretic rewrite system; confluence and normalisation are proved; and the higher groupoid structure is derived from the rewrite theory.

        Finally, the Path vs Eq distinction clarifies how to interpret “failure of UIP” in this setting. The book does not claim that equality in mathematics is ambiguous or that proof-irrelevance is wrong. It claims that there is a mathematically meaningful refinement of equality—equality with reason—that is naturally present in proof theory and rewriting, and that this refinement has genuinely nontrivial homotopical content.

        For the intended audience of mathematicians (algebraists, topologists, type theorists, proof theorists), this provides a familiar pattern: introduce a refinement that carries structure (Path), study its quotient and coherence (RwEq and higher derivations), and recover classical invariants (fundamental groupoids, loop groups) as derived constructions.

        The proposal is therefore intentionally conservative about foundations and ambitious about structure: conservative in that it keeps ordinary equality as the ambient meta-level notion, and ambitious in that it extracts a rich higher-dimensional algebra from explicitly recorded rewrite data. This combination is what makes the project suitable as a research monograph aimed at a broad mathematical audience.
        """
    )


# -----------------------------
# Main conversion
# -----------------------------

def parse_markdown_into_doc(doc: Document, md: str) -> None:
    lines = md.splitlines()

    buf: list[str] = []
    bullet_buf: list[str] | None = None
    num_buf: list[str] | None = None

    def flush_paragraph_buf() -> None:
        nonlocal buf
        if buf:
            add_paragraphs_from_block(doc, "\n".join(buf))
            buf = []

    def flush_lists() -> None:
        nonlocal bullet_buf, num_buf
        if bullet_buf is not None:
            add_bullets(doc, bullet_buf)
            bullet_buf = None
        if num_buf is not None:
            add_numbered(doc, num_buf)
            num_buf = None

    inserted_ch1_expansion = False
    in_ch1 = False

    for raw in lines:
        line = raw.rstrip("\n")

        # Skip top title area; we already have a title page.
        if line.startswith("# Book Proposal"):
            continue

        if line.strip() == "---":
            continue

        # Headings
        m = re.match(r"^(#{2,4})\s+(.*)$", line)
        if m:
            hashes, title = m.group(1), m.group(2)
            title = strip_inline_md(title)

            # Before any new heading, flush accumulated paragraphs/lists.
            flush_paragraph_buf()
            flush_lists()

            # Insert expanded Chapter 1 immediately before Chapter 2 begins.
            if title.startswith("Chapter 2") and in_ch1 and (not inserted_ch1_expansion):
                add_ch1_expansion(doc)
                inserted_ch1_expansion = True
                in_ch1 = False

            # Determine heading level mapping.
            if hashes == "##":
                # Major proposal sections: use Heading 1
                doc.add_heading(title, level=1)
            elif hashes == "###":
                # Parts should be Heading 1; other ### use Heading 3
                if title.startswith("Part "):
                    doc.add_heading(title, level=1)
                else:
                    doc.add_heading(title, level=3)
            elif hashes == "####":
                # Chapters / Appendices: Heading 2
                doc.add_heading(title, level=2)
                if title.startswith("Chapter 1"):
                    in_ch1 = True
                elif title.startswith("Chapter "):
                    in_ch1 = False
            continue

        # Bullets
        bm = re.match(r"^\s*[-•]\s+(.*)$", line)
        if bm:
            flush_paragraph_buf()
            if num_buf is not None:
                flush_lists()
            if bullet_buf is None:
                bullet_buf = []
            bullet_buf.append(bm.group(1))
            continue

        # Numbered lists like "1." "2." etc
        nm = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if nm:
            flush_paragraph_buf()
            if bullet_buf is not None:
                flush_lists()
            if num_buf is None:
                num_buf = []
            num_buf.append(nm.group(1))
            continue

        # Blank line
        if not line.strip():
            flush_paragraph_buf()
            flush_lists()
            continue

        # Regular text line
        buf.append(line)

    flush_paragraph_buf()
    flush_lists()

    # If Chapter 1 was last or Chapter 2 heading wasn't seen for some reason, still insert.
    if in_ch1 and not inserted_ch1_expansion:
        add_ch1_expansion(doc)


def main() -> None:
    md = IN_MD.read_text(encoding="utf-8")

    doc = Document()
    set_default_styles(doc)
    add_title_page(doc)

    parse_markdown_into_doc(doc, md)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(str(OUT_DOCX))


if __name__ == "__main__":
    main()
